"""ResumeData → DOCX 渲染器。

使用 python-docx 生成简洁专业风格的 Word 文档，
单栏布局、标准字体、清晰层级，适合 ATS 系统解析。
"""

from __future__ import annotations

import io
import logging
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from resume_agent.models.resume_data import ResumeData

log = logging.getLogger(__name__)


def render_resume_data_to_docx(data: ResumeData, template: str = "professional") -> bytes:
    """将 ResumeData 渲染为 DOCX 字节数据。

    Args:
        data: 简历结构化数据
        template: 模板名称（docx 输出统一使用简洁专业风格，模板名仅影响配色细节）

    Returns:
        DOCX 文件字节数据
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as exc:
        raise ImportError("python-docx 未安装，请执行: pip install python-docx") from exc

    doc = Document()

    # 页面边距
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # 默认字体
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.space_before = Pt(0)

    # 模板配色
    colors = {
        "professional": RGBColor(0x2D, 0x3A, 0x4A),  # 深蓝灰
        "academic": RGBColor(0x1A, 0x1A, 0x2E),       # 深色
        "creative": RGBColor(0x4A, 0x2D, 0x6A),        # 紫色
        "minimal": RGBColor(0x18, 0x18, 0x1B),          # 纯黑
        "elegant": RGBColor(0x1A, 0x23, 0x32),          # 深蓝
        "tech": RGBColor(0x11, 0x18, 0x27),             # 深色终端
        "compact": RGBColor(0x1F, 0x29, 0x37),          # 暗灰
    }
    heading_color = colors.get(template, colors["professional"])

    # ---------- 姓名 ----------
    if data.name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(data.name)
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = heading_color
        p.paragraph_format.space_after = Pt(2)

    # ---------- 联系方式 ----------
    contact_parts: list[str] = []
    if data.contact.email:
        contact_parts.append(data.contact.email)
    if data.contact.phone:
        contact_parts.append(data.contact.phone)
    if data.contact.location:
        contact_parts.append(data.contact.location)
    if data.contact.website:
        contact_parts.append(data.contact.website)
    if data.contact.linkedin:
        contact_parts.append(data.contact.linkedin)
    if data.contact.wechat:
        contact_parts.append(f"微信: {data.contact.wechat}")
    if not contact_parts and data.contact.raw_text:
        contact_parts.append(data.contact.raw_text)

    if contact_parts:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(" | ".join(contact_parts))
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        p.paragraph_format.space_after = Pt(6)

    # 添加分隔线
    _add_horizontal_rule(doc)

    # ---------- 按 section_order 渲染各章节 ----------
    default_order = ["summary", "experience", "education", "skills", "projects"]
    section_order = data.section_order if data.section_order else default_order

    section_builders = {
        "summary": _build_summary,
        "experience": _build_experience,
        "education": _build_education,
        "skills": _build_skills,
        "projects": _build_projects,
    }

    rendered_keys: set[str] = set()
    for key in section_order:
        if key in rendered_keys:
            continue
        if key in section_builders:
            section_builders[key](doc, data, heading_color)
            rendered_keys.add(key)

    # 渲染 section_order 中未包含但有内容的章节
    for key in default_order:
        if key not in rendered_keys and key in section_builders:
            section_builders[key](doc, data, heading_color)

    # 输出字节
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_horizontal_rule(doc: "Document") -> None:
    """添加水平分隔线。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    # 使用底部边框模拟分隔线
    from docx.oxml.ns import qn
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single",
        qn("w:sz"): "6",
        qn("w:space"): "1",
        qn("w:color"): "CCCCCC",
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_section_heading(doc: "Document", title: str, color) -> None:
    """添加章节标题。"""
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)

    # 标题下方添加细线
    from docx.oxml.ns import qn
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single",
        qn("w:sz"): "4",
        qn("w:space"): "1",
        qn("w:color"): "DDDDDD",
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def _build_summary(doc: "Document", data: "ResumeData", color) -> None:
    """构建个人简介章节。"""
    if not data.summary:
        return
    _add_section_heading(doc, "个人简介", color)
    p = doc.add_paragraph(data.summary)
    p.paragraph_format.space_after = Pt(6)


def _build_experience(doc: "Document", data: "ResumeData", color) -> None:
    """构建工作经历章节。"""
    if not data.experience:
        return
    _add_section_heading(doc, "工作经历", color)
    for exp in data.experience:
        # 职位 + 公司 + 时间
        p = doc.add_paragraph()
        run = p.add_run(exp.title)
        run.bold = True
        run.font.size = Pt(11)
        if exp.company:
            run = p.add_run(f"  |  {exp.company}")
            run.font.size = Pt(10.5)
        p.paragraph_format.space_after = Pt(1)

        # 时间
        if exp.period:
            p_time = doc.add_paragraph()
            run = p_time.add_run(exp.period)
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            run.italic = True
            p_time.paragraph_format.space_after = Pt(2)

        # 成果
        for h in exp.highlights:
            p = doc.add_paragraph(style="List Bullet")
            p.text = h
            p.paragraph_format.space_after = Pt(1)
            for run in p.runs:
                run.font.size = Pt(10)

        # 间距
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.space_before = Pt(0)


def _build_education(doc: "Document", data: "ResumeData", color) -> None:
    """构建教育背景章节。"""
    if not data.education:
        return
    _add_section_heading(doc, "教育背景", color)
    for edu in data.education:
        # 学位 + 专业 + 学校
        p = doc.add_paragraph()
        parts = []
        if edu.degree:
            parts.append(edu.degree)
        if edu.major:
            parts.append(edu.major)
        if parts:
            run = p.add_run(" - ".join(parts))
            run.bold = True
            run.font.size = Pt(11)
        if edu.school:
            run = p.add_run(f"  |  {edu.school}")
            run.font.size = Pt(10.5)
        p.paragraph_format.space_after = Pt(1)

        # 时间
        if edu.period:
            p_time = doc.add_paragraph()
            run = p_time.add_run(edu.period)
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            run.italic = True
            p_time.paragraph_format.space_after = Pt(2)

        # 成就
        for a in edu.achievements:
            p = doc.add_paragraph(style="List Bullet")
            p.text = a
            p.paragraph_format.space_after = Pt(1)
            for run in p.runs:
                run.font.size = Pt(10)

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)


def _build_skills(doc: "Document", data: "ResumeData", color) -> None:
    """构建专业技能章节。"""
    if not data.skills:
        return
    _add_section_heading(doc, "专业技能", color)
    for cat in data.skills:
        p = doc.add_paragraph()
        run = p.add_run(f"{cat.category}：")
        run.bold = True
        run.font.size = Pt(10.5)
        run = p.add_run("、".join(cat.skills))
        run.font.size = Pt(10.5)
        p.paragraph_format.space_after = Pt(2)


def _build_projects(doc: "Document", data: "ResumeData", color) -> None:
    """构建项目经历章节。"""
    if not data.projects:
        return
    _add_section_heading(doc, "项目经历", color)
    for proj in data.projects:
        # 项目名 + 角色 + 时间
        p = doc.add_paragraph()
        run = p.add_run(proj.name)
        run.bold = True
        run.font.size = Pt(11)
        if proj.role:
            run = p.add_run(f"  |  {proj.role}")
            run.font.size = Pt(10.5)
        p.paragraph_format.space_after = Pt(1)

        if proj.period:
            p_time = doc.add_paragraph()
            run = p_time.add_run(proj.period)
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            run.italic = True
            p_time.paragraph_format.space_after = Pt(2)

        if proj.description:
            p = doc.add_paragraph(style="List Bullet")
            p.text = f"项目描述：{proj.description}"
            p.paragraph_format.space_after = Pt(1)
            for run in p.runs:
                run.font.size = Pt(10)

        for c in proj.contributions:
            p = doc.add_paragraph(style="List Bullet")
            p.text = c
            p.paragraph_format.space_after = Pt(1)
            for run in p.runs:
                run.font.size = Pt(10)

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
